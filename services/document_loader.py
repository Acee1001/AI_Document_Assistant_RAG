"""Utilities for loading and parsing uploaded documents."""

import os
import shutil
import tempfile
import zipfile


def load_pdf(file_path: str) -> str:
    """Load text content from a PDF file."""
    import fitz

    with fitz.open(file_path) as doc:
        return "\n".join(page.get_text() for page in doc).strip()


def load_txt(file_path: str) -> str:
    """Load text content from a UTF-8 text file."""
    with open(file_path, "r", encoding="utf-8", errors="ignore") as file:
        return file.read().strip()


def load_docx(file_path: str) -> str:
    """Load text content from DOCX paragraphs and table cells."""
    from docx import Document

    doc = Document(file_path)
    text: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text.append(cell.text)
    return "\n".join(text).strip()


def load_zip(file_path: str) -> str:
    """Load and combine text from all supported files inside a ZIP archive."""
    text_parts: list[str] = []
    tmp_dir = tempfile.mkdtemp(prefix="rag_zip_")
    try:
        with zipfile.ZipFile(file_path, "r") as archive:
            archive.extractall(tmp_dir)
        for root, _, files in os.walk(tmp_dir):
            for fname in files:
                fpath = os.path.join(root, fname)
                ext = os.path.splitext(fname)[1].lower()
                try:
                    if ext == ".pdf":
                        content = load_pdf(fpath)
                    elif ext == ".txt":
                        content = load_txt(fpath)
                    elif ext == ".docx":
                        content = load_docx(fpath)
                    else:
                        continue
                    text_parts.append(f"=== File: {fname} ===\n{content}")
                except Exception as exc:
                    text_parts.append(f"=== File: {fname} - Error: {exc} ===")
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)
    return "\n\n".join(text_parts).strip()


def load_document(file_path: str) -> str:
    """Route a file path to the matching loader by extension."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return load_pdf(file_path)
    if ext == ".txt":
        return load_txt(file_path)
    if ext == ".docx":
        return load_docx(file_path)
    if ext == ".zip":
        return load_zip(file_path)
    raise ValueError(f"Unsupported file type: {ext}")
